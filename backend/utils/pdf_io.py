"""PDF + DOCX + image ingestion helpers.

Three input formats produce the same shape:
    {
        "page_count": int,
        "pages": [{"page_number": int, "text": str, "has_tables": bool}, ...],
        "is_scanned": bool,   # True if PDF has no embedded text / image input
    }

PDFs use pdfplumber for text + pdf2image for raster pages.
DOCX uses python-docx; the whole document is treated as one logical page.
Images go straight to OCR (treated as one-page scanned documents).
"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ─── PDF ─────────────────────────────────────────────────────────────────


def parse_pdf(file_path: str) -> dict:
    """Return {page_count, pages, is_scanned} for a PDF."""
    try:
        import pdfplumber
    except ImportError:
        return _error("pdfplumber not installed")

    if not os.path.exists(file_path):
        return _error(f"File not found: {file_path}")

    try:
        with pdfplumber.open(file_path) as pdf:
            pages_out: list[dict] = []
            embedded_text_chars = 0
            for i, page in enumerate(pdf.pages):
                text = (page.extract_text() or "").strip()
                tables = page.extract_tables() or []
                embedded_text_chars += len(text)
                pages_out.append({
                    "page_number": i + 1,
                    "text": text,
                    "has_tables": bool(tables),
                })
            return {
                "page_count": len(pages_out),
                "pages": pages_out,
                "is_scanned": embedded_text_chars < 100 * len(pages_out),
            }
    except Exception as exc:
        return _error(f"pdfplumber failed: {type(exc).__name__}: {exc}")


def rasterise_pdf(file_path: str, output_dir: str, dpi: int = 200) -> list[str]:
    """Render each PDF page to a PNG. Returns list of file paths."""
    try:
        from pdf2image import convert_from_path
    except ImportError:
        logger.warning("pdf2image not installed; cannot rasterise PDF")
        return []

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    stem = Path(file_path).stem
    try:
        images = convert_from_path(file_path, dpi=dpi)
    except Exception as exc:
        logger.error("pdf2image failed: %s", exc)
        return []

    paths: list[str] = []
    for i, img in enumerate(images):
        out = os.path.join(output_dir, f"{stem}_page_{i + 1}.png")
        img.save(out, format="PNG", optimize=True)
        paths.append(out)
    return paths


# ─── DOCX ────────────────────────────────────────────────────────────────


def parse_docx(file_path: str) -> dict:
    if not os.path.exists(file_path):
        return _error(f"File not found: {file_path}")
    try:
        from docx import Document
    except ImportError:
        return _error("python-docx not installed")
    try:
        doc = Document(file_path)
    except Exception as exc:
        return _error(f"python-docx failed: {type(exc).__name__}: {exc}")

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text)
    has_tables = bool(doc.tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").replace("\n", " ").strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return {
        "page_count": 1,
        "pages": [{"page_number": 1, "text": "\n".join(parts), "has_tables": has_tables}],
        "is_scanned": False,
    }


# ─── Images ──────────────────────────────────────────────────────────────


def parse_image(file_path: str) -> dict:
    if not os.path.exists(file_path):
        return _error(f"File not found: {file_path}")
    return {
        "page_count": 1,
        "pages": [{"page_number": 1, "text": "", "has_tables": False}],
        "is_scanned": True,
    }


def copy_image_as_page(image_path: str, output_dir: str) -> list[str]:
    """Normalise an uploaded image to a PNG so OpenCV/Tesseract can read it."""
    try:
        from PIL import Image
    except ImportError:
        return []
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    stem = Path(image_path).stem
    out = os.path.join(output_dir, f"{stem}_page_1.png")
    try:
        with Image.open(image_path) as img:
            if getattr(img, "is_animated", False):
                try:
                    img.seek(0)
                except EOFError:
                    pass
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            img.save(out, format="PNG", optimize=True)
        return [out]
    except Exception as exc:
        logger.error("Pillow failed to convert %s: %s", image_path, exc)
        return []


# ─── Helpers ─────────────────────────────────────────────────────────────


def _error(message: str) -> dict:
    return {
        "error": True,
        "message": message,
        "page_count": 0,
        "pages": [],
        "is_scanned": False,
    }
